from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from typing import List
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os
import tempfile

app = FastAPI()

# CORS - Permite llamadas desde cualquier origen
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # En producción: específica tu dominio
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ===== MODELOS =====
class Terreno(BaseModel):
    valor_terreno_propio: float
    metros_terreno_propio: float
    valor_terreno_comun: float = 0.0
    metros_terreno_comun: float = 0.0

class Construccion(BaseModel):
    valor_construccion_propia: float
    metros_construccion_propia: float
    valor_construccion_comun: float = 0.0
    metros_construccion_comun: float = 0.0

class Impuesto(BaseModel):
    impuesto_predial: float
    cantidad_con_letra: str
    recargos: float = 0.0
    gastos_ejecucion: float = 0.0
    subsidios: float = 0.0
    ultimo_periodo_pagado: str = "2024"

class Predio(BaseModel):
    clave_catastral: str = Field(..., pattern=r'^\d{3}-\d{2}-\d{3}-\d{2}-\d{2}-[A-Z0-9]+

# ===== ENDPOINTS =====
@app.get("/")
def root():
    return {"status": "ok", "message": "API funcionando correctamente"}

@app.get("/api")
def api_root():
    return {"status": "ok", "endpoints": ["/api/generar-docx"]}

@app.post("/api/generar-docx")
async def generar_docx(file: UploadFile = File(...)):
    """
    Recibe un archivo JSON y genera un documento DOCX
    """
    try:
        # Validar que sea JSON
        if not file.filename.endswith('.json'):
            raise HTTPException(
                status_code=400, 
                detail="Solo se permiten archivos JSON"
            )
        
        # Leer y validar contenido
        content = await file.read()
        data = json.loads(content)
        validated_data = DatosRequest(**data)
        
        # Buscar plantilla Word
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', '1785-003.docx')
        
        if os.path.exists(template_path):
            # Cargar plantilla existente
            doc = Document(template_path)
            print(f"✅ Plantilla cargada: {template_path}")
            
            # Reemplazar placeholders en la plantilla
            # Basado en el formato: {{ p.variable }}
            replacements = {}
            
            # Procesar cada predio (si hay múltiples, usa el primero)
            if len(validated_data.predio) >= 1:
                predio = validated_data.predio[0]
                
                # Calcular valores derivados
                terreno_total_metros = predio.terreno.metros_terreno_propio + predio.terreno.metros_terreno_comun
                terreno_total_valor = predio.terreno.valor_terreno_propio + predio.terreno.valor_terreno_comun
                construccion_total_metros = predio.construccion.metros_construccion_propia + predio.construccion.metros_construccion_comun
                construccion_total_valor = predio.construccion.valor_construccion_propia + predio.construccion.valor_construccion_comun
                
                # Calcular impuestos (usar valores del JSON si existen)
                impuesto_predial = predio.impuesto.impuesto_predial
                recargos = predio.impuesto.recargos if predio.impuesto.recargos > 0 else impuesto_predial * 0.02
                gastos_ejecucion = predio.impuesto.gastos_ejecucion if predio.impuesto.gastos_ejecucion > 0 else impuesto_predial * 0.01
                subsidio = predio.impuesto.subsidios if predio.impuesto.subsidios > 0 else 0
                impuesto_suma = impuesto_predial + recargos + gastos_ejecucion
                adeudo_total = impuesto_suma - subsidio
                
                from datetime import datetime
                fecha_actual = datetime.now().strftime("%d/%m/%Y")
                
                replacements = {
                    # Documento
                    'p.documento.fecha_actual': fecha_actual,
                    'p.clave_catastral': predio.clave_catastral,
                    'p.folio': str(predio.folio),
                    
                    # Contribuyente
                    'p.contribuente': predio.contribuyente,
                    'p.direccion': predio.direccion,
                    
                    # Terreno
                    'p.terreno.metros_terreno_propio': f"{predio.terreno.metros_terreno_propio:,.2f}",
                    'p.terreno.valor_catastral_orig': f"{predio.terreno.valor_terreno_propio:,.2f}",
                    
                    # Construcción
                    'p.construccion.metros_construccion_propia': f"{predio.construccion.metros_construccion_propia:,.2f}",
                    'p.construccion.valor_catastral_orig': f"{predio.construccion.valor_construccion_propia:,.2f}",
                    
                    # Impuestos
                    'p.impuesto_ultimo_periodo_pagado': predio.impuesto.ultimo_periodo_pagado,
                    'p.impuesto.impuesto_predial': f"{impuesto_predial:,.2f}",
                    'p.impuesto.recargos': f"{recargos:,.2f}",
                    'p.impuesto.gastos_ejecucion': f"{gastos_ejecucion:,.2f}",
                    'p.impuesto.suma': f"{impuesto_suma:,.2f}",
                    'p.impuesto.subsidios': f"{subsidio:,.2f}",
                    'p.impuesto.adeudo_a_pagar': f"{adeudo_total:,.2f}",
                    'p.impuesto.texto_monto_con_letra': predio.impuesto.cantidad_con_letra,
                    
                    # Fecha fin de vigencia
                    'p.documento.fecha_fin_vigencia': datetime.now().strftime("%d/%m/%Y")
                }
            
            # Función para reemplazar en texto manteniendo formato
            def replace_in_runs(paragraph, old_text, new_text):
                """Reemplaza texto manteniendo el formato de los runs"""
                if old_text in paragraph.text:
                    # Reconstruir el texto completo
                    full_text = paragraph.text
                    full_text = full_text.replace(old_text, new_text)
                    
                    # Limpiar runs existentes
                    for run in paragraph.runs:
                        run.text = ''
                    
                    # Agregar nuevo texto en el primer run
                    if paragraph.runs:
                        paragraph.runs[0].text = full_text
                    else:
                        paragraph.add_run(full_text)
            
            # Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    placeholder = f'{{{{ {key} }}}}'
                    if placeholder in paragraph.text:
                        replace_in_runs(paragraph, placeholder, str(value))
            
            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, value in replacements.items():
                                placeholder = f'{{{{ {key} }}}}'
                                if placeholder in para.text:
                                    replace_in_runs(para, placeholder, str(value))
        else:
            # Si no existe la plantilla, crear documento básico
            print(f"⚠️ Plantilla no encontrada: {template_path}")
            doc = Document()
            doc.add_heading('Documento Generado Automáticamente', 0)
            doc.add_paragraph(f"Archivo: {validated_data.archivo}")
            doc.add_paragraph(f"Total de predios: {len(validated_data.predio)}")
            doc.add_paragraph("")
            
            for idx, predio in enumerate(validated_data.predio, 1):
                doc.add_heading(f'Predio #{idx}', level=1)
                doc.add_paragraph(f"Clave Catastral: {predio.clave_catastral}")
                doc.add_paragraph(f"Folio: {predio.folio}")
                doc.add_paragraph(f"Dirección: {predio.direccion}")
                doc.add_paragraph(f"Contribuyente: {predio.contribuyente}")
                doc.add_paragraph(f"Impuesto Predial: ${predio.impuesto.impuesto_predial:,.2f}")
                doc.add_paragraph("")
        
        # Guardar en archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Crear generador para enviar archivo
        def file_iterator():
            with open(tmp_path, 'rb') as f:
                yield from f
            # Eliminar archivo temporal después de enviar
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        # Retornar archivo
        return StreamingResponse(
            file_iterator(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={validated_data.archivo}.docx",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=422, 
            detail=f"JSON inválido: {str(e)}"
        )
    except ValueError as e:
        raise HTTPException(
            status_code=422, 
            detail=f"Error de validación: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Error interno del servidor: {str(e)}"
        )

# Vercel necesita que exportes 'app'
# No agregues 'handler' ni nada más al final)
    folio: int
    direccion: str
    contribuyente: str
    terreno: Terreno
    construccion: Construccion
    impuesto: Impuesto

class DatosRequest(BaseModel):
    archivo: str
    predio: List[Predio]

# ===== ENDPOINTS =====
@app.get("/")
def root():
    return {"status": "ok", "message": "API funcionando correctamente"}

@app.get("/api")
def api_root():
    return {"status": "ok", "endpoints": ["/api/generar-docx"]}

@app.post("/api/generar-docx")
async def generar_docx(file: UploadFile = File(...)):
    """
    Recibe un archivo JSON y genera un documento DOCX
    """
    try:
        # Validar que sea JSON
        if not file.filename.endswith('.json'):
            raise HTTPException(
                status_code=400, 
                detail="Solo se permiten archivos JSON"
            )
        
        # Leer y validar contenido
        content = await file.read()
        data = json.loads(content)
        validated_data = DatosRequest(**data)
        
        # Buscar plantilla Word
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', '1785-003.docx')
        
        if os.path.exists(template_path):
            # Cargar plantilla existente
            doc = Document(template_path)
            print(f"✅ Plantilla cargada: {template_path}")
            
            # Reemplazar placeholders en la plantilla
            # Basado en el formato: {{ p.variable }}
            replacements = {}
            
            # Procesar cada predio (si hay múltiples, usa el primero)
            if len(validated_data.predio) >= 1:
                predio = validated_data.predio[0]
                
                # Calcular valores derivados
                terreno_total_metros = predio.terreno.metros_terreno_propio + predio.terreno.metros_terreno_comun
                terreno_total_valor = predio.terreno.valor_terreno_propio + predio.terreno.valor_terreno_comun
                construccion_total_metros = predio.construccion.metros_construccion_propia + predio.construccion.metros_construccion_comun
                construccion_total_valor = predio.construccion.valor_construccion_propia + predio.construccion.valor_construccion_comun
                
                # Calcular impuestos
                impuesto_predial = predio.impuesto.impuesto_predial
                recargos = impuesto_predial * 0.02  # 2% ejemplo
                gastos_ejecucion = impuesto_predial * 0.01  # 1% ejemplo
                impuesto_suma = impuesto_predial + recargos + gastos_ejecucion
                subsidio = impuesto_suma * 0.1  # 10% ejemplo
                adeudo_total = impuesto_suma - subsidio
                
                from datetime import datetime
                fecha_actual = datetime.now().strftime("%d/%m/%Y")
                
                replacements = {
                    # Documento
                    'p.documento.fecha_actual': fecha_actual,
                    'p.clave_catastral': predio.clave_catastral,
                    'p.folio': str(predio.folio),
                    
                    # Contribuyente
                    'p.contribuente': predio.contribuyente,
                    'p.direccion': predio.direccion,
                    
                    # Terreno
                    'p.terreno.metros_terreno_propio': f"{predio.terreno.metros_terreno_propio:,.2f}",
                    'p.terreno.valor_catastral_orig': f"{predio.terreno.valor_terreno_propio:,.2f}",
                    
                    # Construcción
                    'p.construccion.metros_construccion_propia': f"{predio.construccion.metros_construccion_propia:,.2f}",
                    'p.construccion.valor_catastral_orig': f"{predio.construccion.valor_construccion_propia:,.2f}",
                    
                    # Impuestos
                    'p.impuesto_ultimo_periodo_pagado': "2024",  # Ajustar según tu lógica
                    'p.impuesto.impuesto_predial': f"{impuesto_predial:,.2f}",
                    'p.impuesto.recargos': f"{recargos:,.2f}",
                    'p.impuesto.gastos_ejecucion': f"{gastos_ejecucion:,.2f}",
                    'p.impuesto.suma': f"{impuesto_suma:,.2f}",
                    'p.impuesto.subsidios': f"{subsidio:,.2f}",
                    'p.impuesto.adeudo_a_pagar': f"{adeudo_total:,.2f}",
                    'p.impuesto.texto_monto_con_letra': predio.impuesto.cantidad_con_letra,
                    
                    # Fecha fin de vigencia
                    'p.documento.fecha_fin_vigencia': datetime.now().strftime("%d/%m/%Y")
                }
            
            # Función para reemplazar en texto manteniendo formato
            def replace_in_runs(paragraph, old_text, new_text):
                """Reemplaza texto manteniendo el formato de los runs"""
                if old_text in paragraph.text:
                    # Reconstruir el texto completo
                    full_text = paragraph.text
                    full_text = full_text.replace(old_text, new_text)
                    
                    # Limpiar runs existentes
                    for run in paragraph.runs:
                        run.text = ''
                    
                    # Agregar nuevo texto en el primer run
                    if paragraph.runs:
                        paragraph.runs[0].text = full_text
                    else:
                        paragraph.add_run(full_text)
            
            # Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    placeholder = f'{{{{ {key} }}}}'
                    if placeholder in paragraph.text:
                        replace_in_runs(paragraph, placeholder, str(value))
            
            # Reemplazar en tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, value in replacements.items():
                                placeholder = f'{{{{ {key} }}}}'
                                if placeholder in para.text:
                                    replace_in_runs(para, placeholder, str(value))
        else:
            # Si no existe la plantilla, crear documento básico
            print(f"⚠️ Plantilla no encontrada: {template_path}")
            doc = Document()
            doc.add_heading('Documento Generado Automáticamente', 0)
            doc.add_paragraph(f"Archivo: {validated_data.archivo}")
            doc.add_paragraph(f"Total de predios: {len(validated_data.predio)}")
            doc.add_paragraph("")
            
            for idx, predio in enumerate(validated_data.predio, 1):
                doc.add_heading(f'Predio #{idx}', level=1)
                doc.add_paragraph(f"Clave Catastral: {predio.clave_catastral}")
                doc.add_paragraph(f"Folio: {predio.folio}")
                doc.add_paragraph(f"Dirección: {predio.direccion}")
                doc.add_paragraph(f"Contribuyente: {predio.contribuyente}")
                doc.add_paragraph(f"Impuesto Predial: ${predio.impuesto.impuesto_predial:,.2f}")
                doc.add_paragraph("")
        
        # Guardar en archivo temporal
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        # Crear generador para enviar archivo
        def file_iterator():
            with open(tmp_path, 'rb') as f:
                yield from f
            # Eliminar archivo temporal después de enviar
            try:
                os.unlink(tmp_path)
            except:
                pass
        
        # Retornar archivo
        return StreamingResponse(
            file_iterator(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={validated_data.archivo}.docx",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=422, 
            detail=f"JSON inválido: {str(e)}"
        )
    except ValueError as e:
        raise HTTPException(
            status_code=422, 
            detail=f"Error de validación: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Error interno del servidor: {str(e)}"
        )

# Vercel necesita que exportes 'app'
# No agregues 'handler' ni nada más al final
